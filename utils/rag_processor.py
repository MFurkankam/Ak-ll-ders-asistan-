from io import BytesIO
import logging
import os
from typing import List
import uuid
import hashlib

import chromadb
from chromadb.utils import embedding_functions
from chromadb.errors import NotFoundError
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_core.documents import Document
from pypdf import PdfReader
from docx import Document as DocxDocument

logger = logging.getLogger(__name__)

class RAGProcessor:
    """RAG işleme sınıfı - yerel ChromaDB vektör veritabanı"""

    def __init__(self, persist_directory: str = "./chroma_db"):
        self.persist_directory = persist_directory

        self.embedding_function = embedding_functions.DefaultEmbeddingFunction()
        self.chroma_client = chromadb.PersistentClient(path=persist_directory)

        chunk_size = int(os.getenv("RAG_CHUNK_SIZE", "1000"))
        chunk_overlap = int(os.getenv("RAG_CHUNK_OVERLAP", "200"))
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=chunk_size,
            chunk_overlap=chunk_overlap,
            length_function=len,
            separators=[
                "\n\n",
                "\n",
                " ",
                "",
            ],
        )


    def get_dynamic_k(self, query: str, sources_count: int = 0) -> int:
        min_k = int(os.getenv("RAG_MIN_K", "4"))
        max_k = int(os.getenv("RAG_MAX_K", "8"))
        words = len((query or "").split())
        if words >= 20:
            base_k = 8
        elif words >= 12:
            base_k = 6
        else:
            base_k = 4
        if sources_count > 0:
            base_k = base_k + min(4, sources_count // 3)
        return max(min_k, min(max_k, base_k))

    def _resolve_tesseract_cmd(self, pytesseract):
        configured = os.getenv("TESSERACT_CMD")
        if configured:
            pytesseract.pytesseract.tesseract_cmd = configured
            return

        candidates = [
            r"C:\Program Files\Tesseract-OCR\tesseract.exe",
            r"C:\Program Files (x86)\Tesseract-OCR\tesseract.exe",
        ]
        for path in candidates:
            if os.path.exists(path):
                pytesseract.pytesseract.tesseract_cmd = path
                return

    def _extract_text_from_pdf_ocr(self, pdf_bytes: bytes) -> str:
        try:
            import fitz
        except ImportError as exc:
            raise Exception("OCR için PyMuPDF (pymupdf) kurulu değil.") from exc

        try:
            import pytesseract
        except ImportError as exc:
            raise Exception("OCR için pytesseract kurulu değil.") from exc

        try:
            from PIL import Image
        except ImportError as exc:
            raise Exception("OCR için Pillow kurulu değil.") from exc

        self._resolve_tesseract_cmd(pytesseract)

        try:
            pytesseract.get_tesseract_version()
        except Exception as exc:
            raise Exception(
                "Tesseract bulunamadı. Lütfen Tesseract OCR kurun."
            ) from exc

        ocr_dpi = int(os.getenv("OCR_DPI", "150"))
        lang = os.getenv("TESSERACT_LANG")
        texts = []

        with fitz.open(stream=pdf_bytes, filetype="pdf") as doc:
            for page in doc:
                pix = page.get_pixmap(dpi=ocr_dpi)
                mode = "RGB" if pix.alpha == 0 else "RGBA"
                image = Image.frombytes(mode, [pix.width, pix.height], pix.samples)
                if mode == "RGBA":
                    image = image.convert("RGB")

                if lang:
                    try:
                        page_text = pytesseract.image_to_string(image, lang=lang)
                    except Exception:
                        page_text = pytesseract.image_to_string(image)
                else:
                    page_text = pytesseract.image_to_string(image)

                if page_text:
                    texts.append(page_text)

        return "\n".join(texts)

    def extract_text_from_pdf(self, pdf_file) -> str:
        """PDF dosyasından metin çıkar. Metin yoksa OCR dener."""
        try:
            pdf_bytes = pdf_file.read()
            if not pdf_bytes:
                return ""

            text_parts = []
            try:
                pdf_reader = PdfReader(BytesIO(pdf_bytes))
                for page in pdf_reader.pages:
                    text_parts.append(page.extract_text() or "")
            except Exception:
                logger.exception("PDF metin cikarma hatasi")
                text_parts = []

            text = "\n".join(text_parts).strip()
            if len(text) >= 50:
                return text

            return self._extract_text_from_pdf_ocr(pdf_bytes)
        except Exception as exc:
            raise Exception(f"PDF okuma hatası: {str(exc)}") from exc

    def extract_text_from_docx(self, docx_file) -> str:
        """DOCX dosyasından metin çıkar"""
        try:
            doc = DocxDocument(docx_file)
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            return text
        except Exception as exc:
            raise Exception(f"DOCX okuma hatası: {str(exc)}") from exc

    def extract_text_from_txt(self, txt_file) -> str:
        """TXT dosyasından metin çıkar"""
        try:
            content = txt_file.read().decode("utf-8")
            return content
        except Exception as exc:
            raise Exception(f"TXT okuma hatası: {str(exc)}") from exc

    def process_document(self, file, filename: str) -> List[Document]:
        """Dosyayı işle ve parçalara ayır"""
        file_extension = filename.lower().split(".")[-1]

        if file_extension == "pdf":
            text = self.extract_text_from_pdf(file)
        elif file_extension == "docx":
            text = self.extract_text_from_docx(file)
        elif file_extension == "txt":
            text = self.extract_text_from_txt(file)
        else:
            raise ValueError(f"Desteklenmeyen dosya türü: {file_extension}")

        chunks = self.text_splitter.split_text(text)

        documents = [
            Document(
                page_content=chunk,
                metadata={"source": filename, "chunk_id": i},
            )
            for i, chunk in enumerate(chunks)
        ]

        return documents

    def add_documents_to_vectorstore(
        self,
        documents: List[Document],
        collection_name: str = "ders_notlari",
    ):
        """Dokümanları vektör veritabanına ekle"""
        try:
            collection = self.chroma_client.get_or_create_collection(
                name=collection_name,
                embedding_function=self.embedding_function,
            )

            texts = [doc.page_content for doc in documents]
            metadatas = [doc.metadata for doc in documents]
            ids = [
                f"{doc.metadata.get('source', 'unknown')}_{doc.metadata.get('chunk_id', '0')}_{hashlib.sha256(doc.page_content.encode('utf-8')).hexdigest()[:12]}"
                for doc in documents
            ]

            collection.upsert(
                documents=texts,
                metadatas=metadatas,
                ids=ids,
            )

            return collection
        except Exception as exc:
            raise Exception(f"Vektör veritabanına ekleme hatası: {str(exc)}") from exc

    def get_collection(self, collection_name: str = "ders_notlari"):
        """Mevcut koleksiyonu al"""
        try:
            collection = self.chroma_client.get_collection(
                name=collection_name,
                embedding_function=self.embedding_function,
            )
            return collection
        except Exception as exc:
            if isinstance(exc, NotFoundError):
                logger.info("Chroma koleksiyon bulunamadi: %s", collection_name)
            else:
                logger.exception("Chroma koleksiyon alma hatasi")
            return None

    def search_documents(
        self,
        query: str,
        k: int = 4,
        collection_name: str = "ders_notlari",
        source_filter: List[str] | None = None,
    ) -> List[Document]:
        """Sorguya göre en ilgili dokümanları bul"""
        collection = self.get_collection(collection_name)
        if collection is None:
            return []

        try:
            where = None
            if source_filter:
                if len(source_filter) == 1:
                    where = {"source": source_filter[0]}
                else:
                    where = {"source": {"$in": source_filter}}

            if where is None:
                results = collection.query(
                    query_texts=[query],
                    n_results=k,
                )
            else:
                results = collection.query(
                    query_texts=[query],
                    n_results=k,
                    where=where,
                )

            docs = []
            if results and "documents" in results and results["documents"]:
                for i, doc_text in enumerate(results["documents"][0]):
                    metadatas = results.get("metadatas")
                    metadata = metadatas[0][i] if metadatas else {}
                    docs.append(Document(page_content=doc_text, metadata=metadata))

            source_label = "all" if not source_filter else ",".join(source_filter)
            logger.info(f"RAG search: k={k} sources={source_label} results={len(docs)}")

            return docs
        except Exception:
            logger.exception("Chroma sorgu hatasi")
            return []

    def get_all_sources(self, collection_name: str = "ders_notlari") -> List[str]:
        """Veritabanındaki tüm kaynak dosyaları listele"""
        collection = self.get_collection(collection_name)
        if collection is None:
            return []

        try:
            all_data = collection.get()

            sources = set()
            if all_data and "metadatas" in all_data:
                for metadata in all_data["metadatas"]:
                    if metadata and "source" in metadata:
                        sources.add(metadata["source"])

            return sorted(list(sources))
        except Exception:
            logger.exception("Chroma kaynak listeleme hatasi")
            return []

    def delete_collection(self, collection_name: str = "ders_notlari"):
        """Koleksiyonu sil"""
        try:
            self.chroma_client.delete_collection(name=collection_name)
            return True
        except Exception:
            logger.exception("Chroma koleksiyon silme hatasi")
            return False
